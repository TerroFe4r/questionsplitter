import os
import random
import sys
import re

from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *

try:
    import docx

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Библиотека python-docx не установлена. DOCX файлы не будут поддерживаться.")


class PersonEditorDialog(QDialog):
    def __init__(self, initial_persons, parent=None):
        super().__init__(parent)
        self.initial_persons = initial_persons.copy()
        self.persons = initial_persons.copy()
        self.setWindowTitle('Редактирование списка людей')
        self.setGeometry(200, 200, 500, 400)
        self.initUI()

    def initUI(self):
        layout = QVBoxLayout(self)

        label = QLabel('Добавьте, удалите или измените имена людей:')
        layout.addWidget(label)

        self.list_widget = QListWidget()
        self.list_widget.setSelectionMode(QListWidget.SingleSelection)
        self.update_list()
        layout.addWidget(self.list_widget)

        edit_panel = QHBoxLayout()

        self.name_edit = QLineEdit()
        self.name_edit.setPlaceholderText('Введите имя')
        edit_panel.addWidget(self.name_edit)

        add_btn = QPushButton('Добавить')
        add_btn.clicked.connect(self.add_person)
        edit_panel.addWidget(add_btn)

        update_btn = QPushButton('Обновить')
        update_btn.clicked.connect(self.update_person)
        edit_panel.addWidget(update_btn)

        remove_btn = QPushButton('Удалить')
        remove_btn.clicked.connect(self.remove_person)
        edit_panel.addWidget(remove_btn)

        layout.addLayout(edit_panel)

        button_box = QDialogButtonBox(
            QDialogButtonBox.Ok | QDialogButtonBox.Cancel | QDialogButtonBox.Reset
        )
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        button_box.button(QDialogButtonBox.Reset).clicked.connect(self.reset)
        layout.addWidget(button_box)

        self.list_widget.itemSelectionChanged.connect(self.on_selection_changed)

    def update_list(self):
        self.list_widget.clear()
        for person in self.persons:
            self.list_widget.addItem(person)

    def on_selection_changed(self):
        selected_items = self.list_widget.selectedItems()
        if selected_items:
            self.name_edit.setText(selected_items[0].text())

    def add_person(self):
        name = self.name_edit.text().strip()
        if name and name not in self.persons:
            self.persons.append(name)
            self.update_list()
            self.name_edit.clear()

    def update_person(self):
        selected_items = self.list_widget.selectedItems()
        if selected_items and self.name_edit.text().strip():
            old_name = selected_items[0].text()
            new_name = self.name_edit.text().strip()
            if new_name and new_name not in self.persons:
                index = self.persons.index(old_name)
                self.persons[index] = new_name
                self.update_list()

    def remove_person(self):
        selected_items = self.list_widget.selectedItems()
        if selected_items:
            name = selected_items[0].text()
            self.persons.remove(name)
            self.update_list()
            self.name_edit.clear()

    def reset(self):
        self.persons = self.initial_persons.copy()
        self.update_list()
        self.name_edit.clear()

    def get_persons(self):
        return self.persons


class QuestionSplitterApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.default_persons = []
        self.persons = self.default_persons.copy()
        self.questions = []
        self.current_file_path = ""
        self.persons_distribution = {}

        global HAS_DOCX
        if not HAS_DOCX:
            print("Предупреждение: библиотека python-docx не установлена. DOCX файлы не будут поддерживаться.")

        self.initUI()

    def extract_question_number(self, question_text):
        """Извлекает номер вопроса из текста"""
        match = re.search(r'^(\d+)[\.\)\-]', question_text.strip())
        if match:
            return int(match.group(1))
        return 0

    def initUI(self):
        self.setWindowTitle('Распределение вопросов')
        self.setGeometry(100, 100, 1100, 900)

        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)

        top_panel1 = QHBoxLayout()

        self.load_btn = QPushButton('Загрузить файл')
        self.load_btn.setIcon(self.style().standardIcon(QStyle.SP_DialogOpenButton))
        self.load_btn.clicked.connect(self.load_file_dialog)
        self.load_btn.setMinimumHeight(40)
        top_panel1.addWidget(self.load_btn)

        self.edit_persons_btn = QPushButton('👥 Редактировать список людей')
        self.edit_persons_btn.clicked.connect(self.edit_persons)
        self.edit_persons_btn.setMinimumHeight(40)
        top_panel1.addWidget(self.edit_persons_btn)

        self.split_btn = QPushButton('⚖ Равномерное распределение')
        self.split_btn.clicked.connect(self.split_questions)
        self.split_btn.setEnabled(False)
        self.split_btn.setMinimumHeight(40)
        top_panel1.addWidget(self.split_btn)

        self.random_split_btn = QPushButton('🎲 Случайное распределение')
        self.random_split_btn.clicked.connect(self.split_questions_randomly)
        self.random_split_btn.setEnabled(False)
        self.random_split_btn.setMinimumHeight(40)
        top_panel1.addWidget(self.random_split_btn)

        self.save_btn = QPushButton('Сохранить результаты')
        self.save_btn.setIcon(self.style().standardIcon(QStyle.SP_DialogSaveButton))
        self.save_btn.clicked.connect(self.save_results)
        self.save_btn.setEnabled(False)
        self.save_btn.setMinimumHeight(40)
        top_panel1.addWidget(self.save_btn)

        top_panel1.addStretch()

        main_layout.addLayout(top_panel1)

        info_panel = QHBoxLayout()

        file_info_group = QGroupBox("Информация о файле")
        file_info_layout = QVBoxLayout()
        self.file_info_label = QLabel('Файл не выбран')
        self.file_info_label.setWordWrap(True)
        file_info_layout.addWidget(self.file_info_label)
        file_info_group.setLayout(file_info_layout)
        info_panel.addWidget(file_info_group)

        persons_info_group = QGroupBox("Информация о распределении")
        persons_info_layout = QVBoxLayout()
        self.persons_info_label = QLabel(f'Людей: {len(self.persons)}')
        self.questions_info_label = QLabel('Вопросов: 0')
        persons_info_layout.addWidget(self.persons_info_label)
        persons_info_layout.addWidget(self.questions_info_label)
        persons_info_group.setLayout(persons_info_layout)
        info_panel.addWidget(persons_info_group)

        main_layout.addLayout(info_panel)

        self.status_label = QLabel('Выберите файл с вопросами (TXT или DOCX)')
        self.status_label.setStyleSheet("padding: 5px; background-color: #f0f0f0; border-radius: 3px;")
        main_layout.addWidget(self.status_label)

        preview_group = QGroupBox("Предпросмотр вопросов")
        preview_layout = QVBoxLayout()

        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        self.preview_text.setMaximumHeight(450)
        preview_layout.addWidget(self.preview_text)

        preview_group.setLayout(preview_layout)
        main_layout.addWidget(preview_group)

        results_group = QGroupBox("Результаты распределения")
        results_layout = QVBoxLayout()

        self.table = QTableWidget()
        self.table.setColumnCount(3)
        self.table.setHorizontalHeaderLabels(['Имя', 'Кол-во вопросов', 'Вопросы'])
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self.table.setAlternatingRowColors(True)
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        results_layout.addWidget(self.table)

        results_group.setLayout(results_layout)
        main_layout.addWidget(results_group)

        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        main_layout.addWidget(self.progress_bar)

        self.update_persons_info()

    def edit_persons(self):
        dialog = PersonEditorDialog(self.persons, self)
        if dialog.exec_() == QDialog.Accepted:
            new_persons = dialog.get_persons()
            if new_persons:
                self.persons = new_persons
                self.update_persons_info()

                if self.questions:
                    reply = QMessageBox.question(
                        self, 'Перераспределить вопросы?',
                        'Хотите перераспределить вопросы с новым списком людей?',
                        QMessageBox.Yes | QMessageBox.No
                    )
                    if reply == QMessageBox.Yes:
                        self.split_questions()

    def update_persons_info(self):
        self.persons_info_label.setText(f'Людей: {len(self.persons)}')
        if len(self.persons) > 0:
            persons_text = ", ".join(self.persons[:5])
            if len(self.persons) > 5:
                persons_text += f" ... (+{len(self.persons) - 5})"
            self.persons_info_label.setToolTip(f"Список: {persons_text}")

    def load_file_dialog(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите файл с вопросами",
            "",
            "Текстовые файлы (*.txt);;Документы Word (*.docx);;Все файлы (*.*)"
        )

        if file_path:
            self.load_file(file_path)

    def load_file(self, file_path):
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        QApplication.processEvents()

        try:
            self.current_file_path = file_path
            file_name = os.path.basename(file_path)
            file_size = os.path.getsize(file_path) / 1024  # KB
            self.file_info_label.setText(f"Файл: {file_name}\nРазмер: {file_size:.1f} KB")

            if file_path.lower().endswith('.docx'):
                self.questions = self.load_docx_file(file_path)
            else:
                self.questions = self.load_txt_file(file_path)

            self.progress_bar.setValue(50)
            QApplication.processEvents()

            self.status_label.setText(f'Загружено {len(self.questions)} вопросов из файла: {file_name}')
            self.questions_info_label.setText(f'Вопросов: {len(self.questions)}')

            self.show_preview()

            self.split_btn.setEnabled(True)
            self.random_split_btn.setEnabled(True)

            self.progress_bar.setValue(100)
            QMessageBox.information(self, 'Успех',
                                    f'Загружено {len(self.questions)} вопросов')

        except Exception as e:
            QMessageBox.critical(self, 'Ошибка', f'Ошибка при загрузке файла:\n{str(e)}')
        finally:
            self.progress_bar.setVisible(False)

    def load_txt_file(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        questions = []
        lines = content.splitlines()

        for line in lines:
            line = line.strip()
            if line:
                questions.append(line)

        return questions

    def load_docx_file(self, file_path):
        """Загружает вопросы из DOCX файла"""
        if not HAS_DOCX:
            QMessageBox.warning(self, 'Ошибка',
                                'Библиотека python-docx не установлена. Установите:\n'
                                'pip install python-docx')
            return []

        try:
            doc = docx.Document(file_path)
            questions = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    questions.append(text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            questions.append(text)

            return questions

        except Exception as e:
            QMessageBox.critical(self, 'Ошибка', f'Ошибка при чтении DOCX файла:\n{str(e)}')
            return []

    def show_preview(self):
        if not self.questions:
            self.preview_text.clear()
            return

        preview_text = f"Всего вопросов: {len(self.questions)}\n\n"
        preview_text += "Первые 25 вопросов (с сохраненной нумерацией):\n"

        for i, question in enumerate(self.questions[:25], 1):
            display_text = question[:100] + ("..." if len(question) > 100 else "")
            preview_text += f"{display_text}\n"

        if len(self.questions) > 25:
            preview_text += f"\n... и еще {len(self.questions) - 25} вопросов"

        self.preview_text.setText(preview_text)

    def split_questions(self):
        if not self.questions:
            QMessageBox.warning(self, 'Предупреждение', 'Сначала загрузите вопросы')
            return

        if not self.persons:
            QMessageBox.warning(self, 'Предупреждение', 'Добавьте хотя бы одного человека')
            return

        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        QApplication.processEvents()

        distribution = {person: [] for person in self.persons}

        num_persons = len(self.persons)
        for_one_person = len(self.questions) // num_persons
        remainder = len(self.questions) % num_persons

        self.progress_bar.setValue(30)
        QApplication.processEvents()

        for i, person in enumerate(self.persons):
            start_idx = for_one_person * i
            end_idx = for_one_person * (i + 1)
            distribution[person] = self.questions[start_idx:end_idx]

        self.progress_bar.setValue(70)
        QApplication.processEvents()

        if remainder > 0:
            start_remainder = for_one_person * num_persons
            remaining_questions = self.questions[start_remainder:]
            for i, question in enumerate(remaining_questions):
                person = self.persons[i % num_persons]
                distribution[person].append(question)

        self.persons_distribution = distribution
        self.progress_bar.setValue(90)
        QApplication.processEvents()

        self.display_results()
        self.save_btn.setEnabled(True)

        self.progress_bar.setValue(100)
        self.progress_bar.setVisible(False)

    def split_questions_randomly(self):
        if not self.questions:
            QMessageBox.warning(self, 'Предупреждение', 'Сначала загрузите вопросы')
            return

        if not self.persons:
            QMessageBox.warning(self, 'Предупреждение', 'Добавьте хотя бы одного человека')
            return

        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        QApplication.processEvents()

        distribution = {person: [] for person in self.persons}

        num_persons = len(self.persons)
        for_one_person = len(self.questions) // num_persons

        shuffled_questions = self.questions.copy()
        random.shuffle(shuffled_questions)

        self.progress_bar.setValue(30)
        QApplication.processEvents()

        for i, question in enumerate(shuffled_questions):
            person = min(distribution.keys(), key=lambda p: len(distribution[p]))
            distribution[person].append(question)
            if i % 10 == 0:
                progress = 30 + (i / len(shuffled_questions)) * 60
                self.progress_bar.setValue(int(progress))
                QApplication.processEvents()

        self.persons_distribution = distribution
        self.progress_bar.setValue(90)
        QApplication.processEvents()

        self.display_results()
        self.save_btn.setEnabled(True)

        self.progress_bar.setValue(100)
        self.progress_bar.setVisible(False)

    def display_results(self):
        self.table.setRowCount(len(self.persons_distribution))

        total_questions = 0
        row = 0

        for person, questions in self.persons_distribution.items():
            name_item = QTableWidgetItem(person)
            name_item.setFlags(name_item.flags() ^ Qt.ItemIsEditable)
            name_item.setTextAlignment(Qt.AlignCenter)
            self.table.setItem(row, 0, name_item)

            count = len(questions)
            count_item = QTableWidgetItem(str(count))
            count_item.setFlags(count_item.flags() ^ Qt.ItemIsEditable)
            count_item.setTextAlignment(Qt.AlignCenter)
            self.table.setItem(row, 1, count_item)

            # Сортируем вопросы по номеру для отображения в таблице
            sorted_questions = sorted(questions, key=lambda q: self.extract_question_number(q))
            questions_text = "\n".join(sorted_questions)
            questions_item = QTableWidgetItem(questions_text)
            questions_item.setFlags(questions_item.flags() ^ Qt.ItemIsEditable)
            self.table.setItem(row, 2, questions_item)

            total_questions += count
            row += 1

        self.table.resizeRowsToContents()

        min_q = min(len(q) for q in self.persons_distribution.values())
        max_q = max(len(q) for q in self.persons_distribution.values())

        distribution_info = f"Распределено {total_questions} вопросов между {len(self.persons)} людьми"
        if min_q != max_q:
            distribution_info += f" (от {min_q} до {max_q} на человека)"
        else:
            distribution_info += f" (по {min_q} на человека)"

        self.questions_info_label.setText(f'Вопросов распределено: {total_questions}')
        self.status_label.setText(distribution_info)

        self.highlight_extremes()

    def highlight_extremes(self):
        min_q = min(len(q) for q in self.persons_distribution.values())
        max_q = max(len(q) for q in self.persons_distribution.values())

        for row in range(self.table.rowCount()):
            person = self.table.item(row, 0).text()
            count = len(self.persons_distribution[person])

            if count == min_q and min_q != max_q:
                for col in range(3):
                    self.table.item(row, col).setBackground(QColor(200, 255, 200))
            elif count == max_q and min_q != max_q:
                for col in range(3):
                    self.table.item(row, col).setBackground(QColor(255, 200, 200))
            else:
                for col in range(3):
                    self.table.item(row, col).setBackground(QColor(255, 255, 255))

    def save_results(self):
        """Сохраняет результаты в файл DOCX с цветным оформлением"""
        if not hasattr(self, 'persons_distribution') or not self.persons_distribution:
            QMessageBox.warning(self, 'Предупреждение', 'Сначала распределите вопросы')
            return

        # Определяем расширение файла и фильтры
        if not self.current_file_path:
            default_name = "results.docx"
            file_filter = "Документы Word (*.docx);;Текстовые файлы (*.txt);;Все файлы (*.*)"
        else:
            file_dir = os.path.dirname(self.current_file_path)
            file_name = os.path.splitext(os.path.basename(self.current_file_path))[0]
            default_name = os.path.join(file_dir, f"{file_name}_results.docx")
            file_filter = "Документы Word (*.docx);;Текстовые файлы (*.txt);;Все файлы (*.*)"

        save_path, selected_filter = QFileDialog.getSaveFileName(
            self,
            "Сохранить результаты",
            default_name,
            file_filter
        )

        if not save_path:
            return

        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        QApplication.processEvents()

        try:
            # Определяем формат сохранения по расширению файла
            if save_path.lower().endswith('.docx'):
                if not HAS_DOCX:
                    QMessageBox.warning(self, 'Предупреждение',
                                        'Для сохранения в DOCX установите библиотеку python-docx\n'
                                        'pip install python-docx\n\n'
                                        'Результаты будут сохранены в TXT формате.')
                    save_path = save_path.replace('.docx', '.txt')
                    self.save_as_txt(save_path)
                else:
                    self.save_as_docx(save_path)
            else:
                self.save_as_txt(save_path)

            self.progress_bar.setValue(100)
            QMessageBox.information(self, 'Успех',
                                    f'Результаты сохранены в файл:\n{save_path}')

        except Exception as e:
            QMessageBox.critical(self, 'Ошибка', f'Ошибка при сохранении:\n{str(e)}')
        finally:
            self.progress_bar.setVisible(False)

    def save_as_docx(self, file_path):
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()

        styles = doc.styles

        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(16)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle_style = styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_font = subtitle_style.font
        subtitle_font.name = 'Arial'
        subtitle_font.size = Pt(12)
        subtitle_font.bold = True
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        name_style = styles.add_style('CustomName', WD_STYLE_TYPE.PARAGRAPH)
        name_font = name_style.font
        name_font.name = 'Arial'
        name_font.size = Pt(12)
        name_font.bold = True

        question_style = styles.add_style('CustomQuestion', WD_STYLE_TYPE.PARAGRAPH)
        question_font = question_style.font
        question_font.name = 'Arial'
        question_font.size = Pt(11)

        colors = [
            RGBColor(0, 112, 192),  # Синий
            RGBColor(237, 125, 49),  # Оранжевый
            RGBColor(112, 173, 71),  # Зеленый
            RGBColor(255, 192, 0),  # Золотой
            RGBColor(155, 0, 211),  # Фиолетовый
            RGBColor(255, 0, 0),  # Красный
            RGBColor(0, 176, 240),  # Голубой
            RGBColor(146, 208, 80),  # Светло-зеленый
            RGBColor(192, 0, 0),  # Темно-красный
            RGBColor(0, 176, 80),  # Изумрудный
            RGBColor(112, 48, 160),  # Пурпурный
            RGBColor(255, 140, 0),  # Темно-оранжевый
        ]

        title = doc.add_paragraph('РЕЗУЛЬТАТЫ РАСПРЕДЕЛЕНИЯ ВОПРОСОВ', style='CustomTitle')

        total_questions = sum(len(q) for q in self.persons_distribution.values())
        subtitle_text = f"Всего вопросов: {total_questions} | Количество людей: {len(self.persons_distribution)}"
        doc.add_paragraph(subtitle_text, style='CustomSubtitle')

        persons_list = list(self.persons_distribution.keys())

        for i, (person, questions) in enumerate(self.persons_distribution.items()):
            doc.add_paragraph('—' * 39)
            color = colors[i % len(colors)]

            name_para = doc.add_paragraph()
            name_run = name_para.add_run(f"{person} [{len(questions)} вопросов]")
            name_run.font.color.rgb = color
            name_run.font.bold = True
            name_run.font.size = Pt(12)

            doc.add_paragraph('—' * 39)

            sorted_questions = sorted(questions, key=lambda q: self.extract_question_number(q))

            for j, question in enumerate(sorted_questions, 1):
                question_para = doc.add_paragraph(style='CustomQuestion')
                text_run = question_para.add_run(question)
                question_para.paragraph_format.left_indent = Inches(0.2)
                question_para.paragraph_format.space_after = Pt(6)

            doc.add_paragraph()

        doc.add_page_break()
        summary_title = doc.add_paragraph('СВОДНАЯ ТАБЛИЦА РАСПРЕДЕЛЕНИЯ', style='CustomTitle')

        table = doc.add_table(rows=len(persons_list) + 1, cols=3)
        table.style = 'Light Shading'

        header_cells = table.rows[0].cells
        header_cells[0].text = '№'
        header_cells[1].text = 'Имя'
        header_cells[2].text = 'Количество вопросов'

        for i, person in enumerate(persons_list, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(i)
            row_cells[1].text = person
            row_cells[2].text = str(len(self.persons_distribution[person]))

            color = colors[(i - 1) % len(colors)]
            run = row_cells[1].paragraphs[0].runs[0]
            run.font.color.rgb = color

        doc.save(file_path)

    def save_as_txt(self, file_path):
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("=" * 60 + "\n")
            f.write(f"РЕЗУЛЬТАТЫ РАСПРЕДЕЛЕНИЯ ВОПРОСОВ\n")
            f.write(f"Всего вопросов: {sum(len(q) for q in self.persons_distribution.values())}\n")
            f.write(f"Количество людей: {len(self.persons_distribution)}\n")
            f.write("=" * 60 + "\n\n")

            for person, questions in self.persons_distribution.items():
                f.write("-" * 39 + "\n")
                f.write(f"{person} [{len(questions)} вопросов]:\n")
                f.write("-" * 39 + "\n")
                # Сортируем вопросы по номеру
                sorted_questions = sorted(questions, key=lambda q: self.extract_question_number(q))
                for question in sorted_questions:
                    f.write(f"{question}\n")
                f.write("\n")


def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    window = QuestionSplitterApp()
    window.show()
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()